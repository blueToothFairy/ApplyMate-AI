from __future__ import annotations

from pathlib import Path
from docx import Document
from app.config import get_settings
from app.store import store


def export_resume_docx(application_id: str, resume_version_id: str | None = None) -> Path:
    settings = get_settings()
    out_dir = Path(settings.temp_file_dir) / "exports"
    out_dir.mkdir(parents=True, exist_ok=True)

    session = store.get(application_id)
    version_id = resume_version_id or session.selected_resume_version_id
    version = next((v for v in session.tailored_resume_versions if v.version_id == version_id), None)
    if not version:
        raise ValueError(f"Resume version not found: {version_id}")

    doc = Document()
    doc.add_heading("Tailored Resume", level=1)
    for line in version.content.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.strip().startswith("-"):
            doc.add_paragraph(line.strip()[1:].strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())

    path = out_dir / f"{application_id}_{version_id}.docx"
    doc.save(path)
    return path
