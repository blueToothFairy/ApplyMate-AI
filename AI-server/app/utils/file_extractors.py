from __future__ import annotations

from io import BytesIO
from pathlib import Path
from docx import Document
from pypdf import PdfReader


def read_text_file(path: str | Path) -> str:
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def read_docx_file(path: str | Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def extract_text_from_file(path: str | Path) -> str:
    path = Path(path)
    return extract_text_from_bytes(path.read_bytes(), path.name, _guess_mime(path.name))


def extract_text_from_bytes(data: bytes, filename: str, mime_type: str = "application/octet-stream") -> str:
    suffix = Path(filename).suffix.lower()
    normalized_mime = (mime_type or "").lower()

    if suffix in [".txt", ".md"] or normalized_mime.startswith("text/"):
        return data.decode("utf-8", errors="ignore").strip()

    if suffix == ".docx" or "wordprocessingml.document" in normalized_mime:
        doc = Document(BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs).strip()

    if suffix == ".pdf" or normalized_mime == "application/pdf":
        reader = PdfReader(BytesIO(data))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n".join(page for page in pages if page).strip()
        if not text:
            raise ValueError("PDF text extraction returned empty text. The PDF may be scanned/image-only.")
        return text

    raise ValueError(f"Unsupported file type for AI-server parsing: filename={filename}, mime_type={mime_type}")


def _guess_mime(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return "application/pdf"
    if suffix == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if suffix in [".txt", ".md"]:
        return "text/plain"
    return "application/octet-stream"
